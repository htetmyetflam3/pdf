#!/usr/bin/env python3
"""Re-save a .docx with python-docx and mark it as a modern Word 2016 file.

Useful when a DOCX was produced by the old hand-written/SAX merge path and
Word describes it as an "old document" (Compatibility Mode) or python-docx
complains about missing package parts.

Usage:
    python tools/repair_docx.py file1.docx [file2.docx ...]
"""

import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("[!] python-docx is not installed.  Install it with:")
    print("    pip install python-docx")
    print("    Alpine Linux: apk add py3-lxml && pip install python-docx")
    sys.exit(2)


def set_compat_mode_16(doc):
    """Set w:compatSetting compatibilityMode to 16 (Word 2016+)."""
    settings = doc.settings.element
    compat = settings.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        settings.append(compat)
    found = False
    for cs in compat.findall(qn("w:compatSetting")):
        if cs.get(qn("w:name")) == "compatibilityMode":
            cs.set(qn("w:val"), "16")
            found = True
            break
    if not found:
        cs = OxmlElement("w:compatSetting")
        cs.set(qn("w:name"), "compatibilityMode")
        cs.set(qn("w:uri"), "http://schemas.microsoft.com/office/word")
        cs.set(qn("w:val"), "16")
        compat.append(cs)


def repair(path):
    path = str(path)
    src = Path(path)
    if not src.exists():
        raise FileNotFoundError(f"DOCX not found: {path}")

    doc = Document(path)
    set_compat_mode_16(doc)
    out = src.with_name(src.stem + ".repaired.docx")
    doc.save(str(out))

    # Verify the repaired file opens.
    check = Document(str(out))
    print(f"[+] Repaired: {path} -> {out}")
    print(f"    paragraphs={len(check.paragraphs):,}, tables={len(check.tables):,}, "
          f"sections={len(check.sections):,}")
    return out


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        return 1
    failed = False
    for path in sys.argv[1:]:
        try:
            repair(path)
        except Exception as e:
            print(f"[!!] Could not repair {path}: {e}")
            failed = True
    return 1 if failed else 0


if __name__ == "__main__":
    sys.exit(main())
