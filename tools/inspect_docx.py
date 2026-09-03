#!/usr/bin/env python3
"""Inspect Word .docx files with python-docx and print paragraph/line counts.

This is useful for the DOCX chunks produced by `module/formatter.py` and for
the file produced by `tools/merge.py`.  In this pipeline one generated text
line is stored as one Word paragraph, so paragraph count is the practical
line count.

Usage:
    python tools/inspect_docx.py file1.docx [file2.docx ...]
"""

import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("[!] python-docx is not installed.  Install it with:")
    print("    pip install python-docx")
    print("    Alpine Linux: apk add py3-lxml && pip install python-docx")
    sys.exit(2)


def inspect(path):
    path = str(path)
    if not Path(path).exists():
        raise FileNotFoundError(f"DOCX not found: {path}")

    doc = Document(path)
    paragraphs = doc.paragraphs
    content = [p for p in paragraphs if p.text.strip()]
    page_breaks = len(doc.element.body.xpath('.//w:br[@w:type="page"]'))
    chars = sum(len(p.text) for p in paragraphs)

    print(f"DOCX: {path}")
    print(f"  paragraphs       : {len(paragraphs):,}   (one generated text line = one paragraph)")
    print(f"  content lines    : {len(content):,}")
    print(f"  blank lines      : {len(paragraphs) - len(content):,}")
    print(f"  page-break marks : {page_breaks:,}")
    print(f"  tables           : {len(doc.tables):,}")
    print(f"  sections         : {len(doc.sections):,}")
    print(f"  characters       : {chars:,}")
    return {
        "paragraphs": len(paragraphs),
        "content_lines": len(content),
        "blank_lines": len(paragraphs) - len(content),
        "page_breaks": page_breaks,
        "tables": len(doc.tables),
        "sections": len(doc.sections),
        "characters": chars,
    }


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        return 1
    failed = False
    for path in sys.argv[1:]:
        try:
            inspect(path)
        except Exception as e:
            print(f"[!!] Could not inspect {path}: {e}")
            failed = True
    return 1 if failed else 0


if __name__ == "__main__":
    sys.exit(main())
