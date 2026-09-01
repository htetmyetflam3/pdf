#!/usr/bin/env python3
"""Merge large TXT/DOCX chunk files and inspect DOCX content (line counts).

DOCX merge is memory-efficient by default:
  * it streams each ``word/document.xml`` through SAX,
  * it writes the merged body to a temporary file instead of holding the
    whole document in RAM,
  * it keeps only the first input's ``w:sectPr`` (Word section properties),
  * it writes the standard OOXML supporting parts and stamps
    ``compatibilityMode=16`` so Word 2016+ doesn't treat the result as an
    old/compatibility-mode document.

python-docx is still used for ``--inspect``, ``--verify``, and the optional
``--use-python-docx`` merge mode (small files / files with images or
hyperlinks that need relationship remapping).

Usage:
    python tools/merge.py [SOURCE_DIR]
    python tools/merge.py [SOURCE_DIR] --verify
    python tools/merge.py [SOURCE_DIR] --use-python-docx     # small files only
    python tools/merge.py --inspect file1.docx file2.docx
"""

import shutil
import sys
import zipfile
from copy import deepcopy
from io import BytesIO, StringIO
from pathlib import Path
from xml.sax import make_parser, ContentHandler
from xml.sax.handler import feature_external_ges, feature_external_pes

try:
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    HAS_PYTHON_DOCX = True
except Exception:
    HAS_PYTHON_DOCX = False

DOCX_FILES = ["c1to700.docx", "c701to1k.docx", "ending.docx"]
TXT_FILES = ["c1to700.txt", "c701to1k.txt", "ending.txt"]

# Minimal modern OOXML parts.  The streaming writer normally copies these from
# the first template; these constants are used when a part is missing.
_STYLES_XML = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    '<w:docDefaults>'
    '<w:rPrDefault><w:rPr>'
    '<w:rFonts w:ascii="Anonymous Pro" w:hAnsi="Anonymous Pro" w:cs="YoeYar-One" w:eastAsia="YoeYar-One"/>'
    '<w:sz w:val="24"/><w:szCs w:val="24"/>'
    '</w:rPr></w:rPrDefault>'
    '<w:pPrDefault><w:pPr><w:spacing w:line="240" w:lineRule="auto"/></w:pPr></w:pPrDefault>'
    '</w:docDefaults>'
    '<w:style w:type="paragraph" w:default="1" w:styleId="Normal"><w:name w:val="Normal"/><w:qFormat/></w:style>'
    '<w:style w:type="character" w:default="1" w:styleId="DefaultParagraphFont"><w:name w:val="Default Paragraph Font"/></w:style>'
    '<w:style w:type="table" w:default="1" w:styleId="TableNormal"><w:name w:val="Normal Table"/><w:qFormat/></w:style>'
    '<w:style w:type="paragraph" w:styleId="Title"><w:name w:val="Title"/><w:basedOn w:val="Normal"/><w:qFormat/>'
    '<w:pPr><w:spacing w:before="0" w:after="120"/><w:jc w:val="left"/></w:pPr>'
    '<w:rPr><w:b/><w:sz w:val="28"/><w:szCs w:val="28"/></w:rPr></w:style>'
    '<w:style w:type="paragraph" w:styleId="Heading1"><w:name w:val="heading 1"/><w:basedOn w:val="Normal"/><w:qFormat/>'
    '<w:pPr><w:keepNext/><w:spacing w:before="240" w:after="120"/><w:jc w:val="left"/></w:pPr>'
    '<w:rPr><w:b/><w:sz w:val="28"/><w:szCs w:val="28"/></w:rPr></w:style>'
    "</w:styles>"
)

_SETTINGS_XML = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<w:settings xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    '<w:zoom w:percent="100"/>'
    '<w:defaultTabStop w:val="720"/>'
    '<w:compat>'
    '<w:compatSetting w:name="compatibilityMode" w:uri="http://schemas.microsoft.com/office/word" w:val="16"/>'
    "</w:compat>"
    "</w:settings>"
)

_WEB_SETTINGS_XML = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<w:webSettings xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    "<w:optimizeForBrowser/>"
    "</w:webSettings>"
)

_FONT_TABLE_XML = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<w:fonts xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    '<w:font w:name="YoeYar-One"><w:charset w:val="00"/><w:family w:val="auto"/><w:pitch w:val="variable"/></w:font>'
    '<w:font w:name="Anonymous Pro"><w:charset w:val="00"/><w:family w:val="auto"/><w:pitch w:val="variable"/></w:font>'
    "</w:fonts>"
)

_CORE_XML = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" '
    'xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" '
    'xmlns:dcmitype="http://purl.org/dc/dcmitype/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">'
    '<dc:title>Merged Myanmar PDF extraction</dc:title>'
    '<dc:creator>pdf-text-extractor</dc:creator>'
    '<cp:lastModifiedBy>pdf-text-extractor</cp:lastModifiedBy>'
    '<dcterms:modified xsi:type="dcterms:W3CDTF">2026-01-01T00:00:00Z</dcterms:modified>'
    "</cp:coreProperties>"
)

_APP_XML = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" '
    'xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">'
    "<Application>pdf-text-extractor</Application>"
    "<AppVersion>16.0000</AppVersion>"
    "</Properties>"
)

_CONTENT_TYPES_XML = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
    '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
    '<Default Extension="xml" ContentType="application/xml"/>'
    '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
    '<Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>'
    '<Override PartName="/word/settings.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.settings+xml"/>'
    '<Override PartName="/word/fontTable.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.fontTable+xml"/>'
    '<Override PartName="/word/webSettings.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.webSettings+xml"/>'
    '<Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>'
    '<Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>'
    "</Types>"
)

_ROOT_RELS_XML = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
    '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
    '<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>'
    '<Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>'
    "</Relationships>"
)

_DOCUMENT_RELS_XML = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
    '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>'
    '<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/settings" Target="settings.xml"/>'
    '<Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/webSettings" Target="webSettings.xml"/>'
    '<Relationship Id="rId4" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/fontTable" Target="fontTable.xml"/>'
    "</Relationships>"
)

_DOC_XML_HEADER = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
    'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
    "<w:body>"
)

_DOC_XML_FOOTER = "</w:body></w:document>"

# Standard part name -> fallback bytes (used when the template lacks a part).
_STANDARD_PARTS = {
    "[Content_Types].xml": _CONTENT_TYPES_XML,
    "_rels/.rels": _ROOT_RELS_XML,
    "word/_rels/document.xml.rels": _DOCUMENT_RELS_XML,
    "word/styles.xml": _STYLES_XML,
    "word/settings.xml": _SETTINGS_XML,
    "word/webSettings.xml": _WEB_SETTINGS_XML,
    "word/fontTable.xml": _FONT_TABLE_XML,
    "docProps/core.xml": _CORE_XML,
    "docProps/app.xml": _APP_XML,
}


# ---------------------------------------------------------------------------
# TXT: stream line by line
# ---------------------------------------------------------------------------

def merge_txt_files(source_dir, output_dir):
    output_path = output_dir / "merged.txt"

    with open(output_path, "w", encoding="utf-8") as outfile:
        for fname in TXT_FILES:
            src = source_dir / fname
            if not src.exists():
                print(f"Warning: {fname} not found, skipping.")
                continue
            print(f"  Streaming {fname} ...")
            with open(src, "r", encoding="utf-8") as infile:
                shutil.copyfileobj(infile, outfile)
                outfile.write("\n")
    print(f"Merged .txt → {output_path}")


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

def _existing_files(source_dir, names):
    return [source_dir / n for n in names if (source_dir / n).exists()]


def patch_settings_compat_16(data: bytes) -> bytes:
    """Ensure Word sees compatibilityMode=16 (Word 2016+)."""
    import re

    if not data:
        return _SETTINGS_XML.encode("utf-8")
    text = data.decode("utf-8", errors="replace")

    if "compatibilityMode" in text:
        pattern = re.compile(
            r'(w:name="compatibilityMode"\s+w:uri="[^"]*"\s+w:val=")\d+(")'
        )
        text = pattern.sub(r"\g<1>16\g<2>", text)
        # Attribute order can vary; also patch any trailing w:val.
        text = re.sub(
            r'(w:name="compatibilityMode"[^>]*?w:val=")\d+(")',
            r"\g<1>16\g<2>",
            text,
        )
        return text.encode("utf-8")

    inject = (
        '<w:compat>'
        '<w:compatSetting w:name="compatibilityMode" '
        'w:uri="http://schemas.microsoft.com/office/word" w:val="16"/>'
        "</w:compat>"
    )
    if "</w:settings>" in text:
        text = text.replace("</w:settings>", inject + "</w:settings>")
    else:
        text += inject
    return text.encode("utf-8")


def _light_docx_check(path):
    """Check required OOXML parts without loading the whole file."""
    missing = []
    with zipfile.ZipFile(str(path)) as z:
        names = set(z.namelist())
    for required in ("[Content_Types].xml", "word/document.xml",
                     "word/_rels/document.xml.rels", "word/settings.xml"):
        if required not in names:
            missing.append(required)
    return missing


def _set_compat_mode_16(doc):
    """Mark a python-docx package as Word 2016+."""
    settings = doc.settings.element
    compat = settings.find(qn("w:compat"))
    if compat is None:
        compat = OxmlElement("w:compat")
        settings.append(compat)
    found = False
    for cs in compat.findall(qn("w:compatSetting")):
        if cs.get(qn("w:name")) == "compatibilityMode":
            cs.set(qn("w:val"), "16")
            found = True
            break
    if not found:
        cs = OxmlElement("w:compatSetting")
        cs.set(qn("w:name"), "compatibilityMode")
        cs.set(qn("w:uri"), "http://schemas.microsoft.com/office/word")
        cs.set(qn("w:val"), "16")
        compat.append(cs)


# ---------------------------------------------------------------------------
# python-docx merge (optional / small files only)
# ---------------------------------------------------------------------------

def _collect_rel_ids(element):
    ids = set()
    for el in element.iter():
        for attr in (qn("r:id"), qn("r:embed"), qn("r:link")):
            val = el.attrib.get(attr)
            if val:
                ids.add(val)
    return ids


def _map_rels_for_element(src_doc_part, dst_doc_part, element):
    mapping = {}
    for old_rid in sorted(_collect_rel_ids(element)):
        rel = src_doc_part.rels.get(old_rid)
        if rel is None:
            continue
        reltype = rel.reltype
        try:
            if rel.is_external and reltype == RT.HYPERLINK:
                mapping[old_rid] = dst_doc_part.relate_to(
                    rel.target_ref, reltype, is_external=True)
            elif not rel.is_external and reltype == RT.IMAGE:
                mapping[old_rid], _ = dst_doc_part.get_or_add_image(
                    BytesIO(rel.target_part.blob))
            else:
                print(f"  [!] Relationship {old_rid!r} ({reltype}) is not copied "
                      "by the merge helper; formatting/objects may be lost.")
        except Exception as e:
            print(f"  [!] Could not copy relationship {old_rid!r}: {e}")
    return mapping


def _remap_rels(element, mapping):
    for el in element.iter():
        for attr in (qn("r:id"), qn("r:embed"), qn("r:link")):
            old_rid = el.attrib.get(attr)
            if old_rid in mapping:
                el.attrib[attr] = mapping[old_rid]


def merge_docx_files_python_docx(source_dir, output_dir, docx_files=None):
    """Merge with python-docx.  Only for small files / files needing images."""
    docx_files = docx_files or _existing_files(source_dir, DOCX_FILES)
    if not docx_files:
        print("No template .docx found.")
        return

    output_path = output_dir / "merged.docx"
    print("  Building merged.docx with python-docx (loads whole files into RAM) ...")

    base = Document(str(docx_files[0]))
    body = base.element.body
    sectPr = body.find(qn("w:sectPr"))
    for child in list(body):
        body.remove(child)
    if sectPr is not None:
        body.append(sectPr)

    for src_path in docx_files:
        src = Document(str(src_path))
        src_body = src.element.body
        copied = 0
        for child in list(src_body):
            if child.tag == qn("w:sectPr"):
                continue
            new_node = deepcopy(child)
            mapping = _map_rels_for_element(src.part, base.part, new_node)
            _remap_rels(new_node, mapping)
            body.append(new_node)
            copied += 1
        print(f"  Merged {src_path.name}: {copied} body elements")

    _set_compat_mode_16(base)
    base.save(str(output_path))
    print(f"Merged .docx → {output_path}")
    return output_path


# ---------------------------------------------------------------------------
# DOCX inspection (python-docx, one file at a time)
# ---------------------------------------------------------------------------

def inspect_docx(path):
    """Print line/paragraph counts for a .docx using python-docx."""
    path = str(path)
    if not Path(path).exists():
        raise FileNotFoundError(f"DOCX not found: {path}")

    doc = Document(path)
    paragraphs = doc.paragraphs
    content = [p for p in paragraphs if p.text.strip()]
    page_breaks = len(doc.element.body.xpath('.//w:br[@w:type="page"]'))
    chars = sum(len(p.text) for p in paragraphs)

    print(f"DOCX: {path}")
    print(f"  paragraphs      : {len(paragraphs):,}   (each generated text line = 1 paragraph)")
    print(f"  content lines    : {len(content):,}")
    print(f"  blank lines      : {len(paragraphs) - len(content):,}")
    print(f"  page-break marks : {page_breaks:,}")
    print(f"  tables           : {len(doc.tables):,}")
    print(f"  sections         : {len(doc.sections):,}")
    print(f"  characters       : {chars:,}")
    return {
        "paragraphs": len(paragraphs),
        "content_lines": len(content),
        "blank_lines": len(paragraphs) - len(content),
        "page_breaks": page_breaks,
        "tables": len(doc.tables),
        "sections": len(doc.sections),
        "characters": chars,
    }


def verify_docx(path):
    try:
        stats = inspect_docx(path)
        print(f"  [OK] python-docx opened {path}")
        return True
    except Exception as e:
        print(f"  [!!] python-docx could NOT open {path}: {e}")
        return False


def check_docx_inputs(source_dir):
    """Lightweight (zip-level) check that required OOXML parts exist."""
    for path in _existing_files(source_dir, DOCX_FILES):
        try:
            missing = _light_docx_check(path)
            if missing:
                print(f"  [!!] input {path.name} is missing parts: {', '.join(missing)}")
            else:
                print(f"  [i] input OK (package parts): {path.name}")
        except Exception as e:
            print(f"  [!!] input is not a valid ZIP/DOCX: {path} ({e})")


# ---------------------------------------------------------------------------
# Memory-efficient streaming DOCX merge (default)
# ---------------------------------------------------------------------------

class BodyExtractor(ContentHandler):
    """Stream body elements, skipping ``w:sectPr`` and tracking rel-bearing attrs."""

    def __init__(self, out_buffer, rel_flags=None):
        self.out = out_buffer
        self.rel_flags = rel_flags if rel_flags is not None else {}
        self.depth = 0
        self.in_body = False
        self.body_closed = False
        self._skip = 0
        self._buf = []

    def _flush_buf(self):
        if self._buf:
            self.out.write("".join(self._buf).encode("utf-8"))
            self._buf.clear()

    def _collect_attrs(self, attrs):
        for k, v in attrs.items():
            if k in ("r:embed", "r:id", "r:link"):
                self.rel_flags[k] = self.rel_flags.get(k, 0) + 1

    def startElement(self, name, attrs):
        if self.body_closed:
            return
        if name == "w:body":
            self.in_body = True
            self.depth = 1
            return
        if not self.in_body:
            return
        if self._skip:
            self._skip += 1
            return
        if name == "w:sectPr":
            # Keep only the template's final section properties.
            self._skip = 1
            return
        self.depth += 1
        self._collect_attrs(attrs)
        tag = ["<", name]
        for k, v in attrs.items():
            tag.append(" ")
            tag.append(k)
            tag.append('="')
            tag.append(v.replace("&", "&amp;").replace('"', "&quot;").replace("<", "&lt;"))
            tag.append('"')
        tag.append(">")
        self._buf.append("".join(tag))

    def endElement(self, name):
        if self._skip:
            self._skip -= 1
            return
        if self.body_closed:
            return
        if name == "w:body":
            self.in_body = False
            self.body_closed = True
            self._flush_buf()
            return
        if not self.in_body:
            return
        self.depth -= 1
        self._buf.append("</")
        self._buf.append(name)
        self._buf.append(">")
        if self.depth == 1:
            self._flush_buf()

    def characters(self, content):
        if not self.in_body or self.body_closed or self._skip:
            return
        self._buf.append(
            content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        )
        if len(self._buf) > 500:
            self._flush_buf()


class SectPrExtractor(ContentHandler):
    """Capture the last ``w:sectPr`` element while streaming (small output)."""

    def __init__(self, out_buffer):
        self.out = out_buffer
        self.in_sect = 0
        self._buf = []
        self._last = None

    def _flush_buf(self):
        if self._buf:
            self._last = "".join(self._buf)
            self._buf = []

    def startElement(self, name, attrs):
        if name == "w:sectPr":
            self._last = None
            self._buf = []
            self.in_sect = 1
            self._buf.append("<w:sectPr")
            for k, v in attrs.items():
                self._buf.append(" ")
                self._buf.append(k)
                self._buf.append('="')
                self._buf.append(v.replace("&", "&amp;").replace('"', "&quot;").replace("<", "&lt;"))
                self._buf.append('"')
            self._buf.append(">")
            return
        if self.in_sect:
            self._buf.append("<")
            self._buf.append(name)
            for k, v in attrs.items():
                self._buf.append(" ")
                self._buf.append(k)
                self._buf.append('="')
                self._buf.append(v.replace("&", "&amp;").replace('"', "&quot;").replace("<", "&lt;"))
                self._buf.append('"')
            self._buf.append(">")
            self.in_sect += 1

    def endElement(self, name):
        if self.in_sect:
            self._buf.append("</")
            self._buf.append(name)
            self._buf.append(">")
            if name == "w:sectPr":
                self._flush_buf()
                self.in_sect = 0
            else:
                self.in_sect -= 1

    def characters(self, content):
        if self.in_sect:
            self._buf.append(
                content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            )

    def result(self):
        self._flush_buf()
        return self._last or ""


class Utf8ChunkDecoder:
    """Feed SAX from a byte stream in chunks, handling split multibyte chars."""

    def __init__(self, byte_stream, chunk_size=65536):
        self.stream = byte_stream
        self.chunk_size = chunk_size
        self._carry = b""

    def __iter__(self):
        while True:
            chunk = self.stream.read(self.chunk_size)
            if not chunk:
                if self._carry:
                    yield self._carry.decode("utf-8", errors="ignore")
                break
            data = self._carry + chunk
            cutoff = len(data)
            self._carry = b""
            for i in range(min(4, len(data)), 0, -1):
                b = data[-i]
                if b < 0x80 or b >= 0xC0:
                    cutoff = len(data) - i + 1
                    try:
                        data[:cutoff].decode("utf-8")
                    except UnicodeDecodeError:
                        continue
                    self._carry = data[cutoff:]
                    break
            else:
                self._carry = data
                continue
            try:
                yield data[:cutoff].decode("utf-8")
            except UnicodeDecodeError:
                continue


def _stream_xml(path, handler, part="word/document.xml"):
    with zipfile.ZipFile(path, "r") as zin:
        if part not in zin.namelist():
            raise KeyError(f"{part} not found in {path}")
        with zin.open(part) as xml_stream:
            parser = make_parser()
            parser.setFeature(feature_external_ges, False)
            parser.setFeature(feature_external_pes, False)
            parser.setContentHandler(handler)
            for text_chunk in Utf8ChunkDecoder(xml_stream, chunk_size=65536):
                parser.feed(text_chunk)
            parser.close()


def extract_body_streaming(docx_path, out_buffer, rel_flags=None):
    handler = BodyExtractor(out_buffer, rel_flags=rel_flags)
    _stream_xml(docx_path, handler)


def extract_sectpr_streaming(docx_path):
    handler = SectPrExtractor(StringIO())
    _stream_xml(docx_path, handler)
    return handler.result()


def _stream_copy_part(zin, zout, name, fallback=None):
    """Copy one ZIP part without loading it into memory.

    If `name` is missing and `fallback` is provided, write the fallback bytes.
    Returns True when the part was written.
    """
    try:
        item = zin.getinfo(name)
    except KeyError:
        if fallback is None:
            return False
        data = fallback if isinstance(fallback, bytes) else fallback.encode("utf-8")
        zout.writestr(name, data)
        return True

    zinfo = zipfile.ZipInfo(name, item.date_time)
    zinfo.compress_type = item.compress_type
    zinfo.create_system = item.create_system
    zinfo.external_attr = item.external_attr
    try:
        with zin.open(item) as src, zout.open(
                zinfo, "w", force_zip64=item.file_size >= zipfile.ZIP64_LIMIT) as dst:
            shutil.copyfileobj(src, dst, 1024 * 1024)
    except Exception:
        with zin.open(item) as src:
            zout.writestr(item, src.read())
    return True


def _write_stream_setting(zin, zout):
    """Write settings.xml patched to Word 2016 compatibility mode."""
    try:
        raw = zin.read("word/settings.xml")
    except KeyError:
        raw = None
    zout.writestr("word/settings.xml", patch_settings_compat_16(raw))


def _template_part_has(zin, name, needles):
    """True if `name` exists in the template and contains every `needle`."""
    try:
        data = zin.read(name).decode("utf-8", errors="replace")
    except KeyError:
        return False
    return all(needle in data for needle in needles)


def _write_content_types(zin, zout):
    """Copy a complete content-types part or fall back to the bundled one."""
    required = (
        "wordprocessingml.document.main",
        "wordprocessingml.styles",
        "wordprocessingml.settings",
        "wordprocessingml.fontTable",
        "wordprocessingml.webSettings",
        "core-properties",
        "extended-properties",
    )
    if _template_part_has(zin, "[Content_Types].xml", required):
        zout.writestr("[Content_Types].xml", zin.read("[Content_Types].xml"))
    else:
        zout.writestr("[Content_Types].xml", _CONTENT_TYPES_XML)


def _write_root_rels(zin, zout):
    """Write a complete package-level relationships part or fall back."""
    required = ("core-properties", "extended-properties")
    if _template_part_has(zin, "_rels/.rels", required):
        zout.writestr("_rels/.rels", zin.read("_rels/.rels"))
    else:
        zout.writestr("_rels/.rels", _ROOT_RELS_XML)


def _write_document_rels(zin, zout):
    """Write a complete document-level relationships part or fall back."""
    required = ("/styles", "relationships/settings", "webSettings", "fontTable")
    if _template_part_has(zin, "word/_rels/document.xml.rels", required):
        zout.writestr("word/_rels/document.xml.rels", zin.read("word/_rels/document.xml.rels"))
    else:
        zout.writestr("word/_rels/document.xml.rels", _DOCUMENT_RELS_XML)


def _write_standard_part(zin, zout, name):
    """Copy the template part if present, else write the bundled fallback."""
    if name in _STANDARD_PARTS:
        fallback = _STANDARD_PARTS[name]
    else:
        fallback = None
    try:
        zout.writestr(name, zin.read(name))
    except KeyError:
        if fallback is not None:
            data = fallback if isinstance(fallback, bytes) else fallback.encode("utf-8")
            zout.writestr(name, data)


def _copy_template_extra_parts(zin, zout, written):
    """Stream all template parts that are not already written / replaced."""
    for item in zin.infolist():
        if item.filename == "word/document.xml":
            continue
        if item.filename in written:
            continue
        _stream_copy_part(zin, zout, item.filename)
        written.add(item.filename)


def merge_docx_files_streaming(source_dir, output_dir, docx_files=None):
    """Memory-efficient merge: SAX-stream bodies, emit a valid Word 2016+ package."""
    docx_files = docx_files or _existing_files(source_dir, DOCX_FILES)
    if not docx_files:
        print("No template .docx found.")
        return

    output_path = output_dir / "merged.docx"
    print("  Building merged.docx with streaming SAX merge (low memory) ...")

    template = docx_files[0]
    temp_body_file = output_dir / ".merged_body_temp.xml"
    temp_docxml = output_dir / ".merged_document.xml"

    rel_flags = {}
    try:
        with open(temp_body_file, "wb") as body_out:
            for src in docx_files:
                print(f"  Streaming body from {src.name} ...")
                extract_body_streaming(src, body_out, rel_flags=rel_flags)

        if rel_flags:
            raise ValueError(
                "The streaming merge only supports text-only DOCX files. "
                f"Detected relationship attributes: {sorted(rel_flags)} "
                "(images, hyperlinks, or objects). Use --use-python-docx for "
                "small files that must preserve images/hyperlinks."
            )

        # Keep the template's final section properties (page size/margins).
        sectpr = extract_sectpr_streaming(template)

        with open(temp_docxml, "wb") as dout:
            dout.write(_DOC_XML_HEADER.encode("utf-8"))
            with open(temp_body_file, "rb") as body_in:
                shutil.copyfileobj(body_in, dout)
            if sectpr:
                dout.write(sectpr.encode("utf-8"))
            dout.write(_DOC_XML_FOOTER.encode("utf-8"))

        with zipfile.ZipFile(str(template), "r") as zin:
            with zipfile.ZipFile(str(output_path), "w", zipfile.ZIP_DEFLATED) as zout:
                written = set()

                # Required package/relationship parts.  Prefer complete template
                # copies; fall back to the bundled modern parts when the template
                # is minimal or missing required overrides/relationships.
                _write_content_types(zin, zout)
                written.add("[Content_Types].xml")
                _write_root_rels(zin, zout)
                written.add("_rels/.rels")
                _write_document_rels(zin, zout)
                written.add("word/_rels/document.xml.rels")
                for name in ("word/styles.xml", "word/webSettings.xml",
                             "word/fontTable.xml",
                             "docProps/core.xml", "docProps/app.xml"):
                    _write_standard_part(zin, zout, name)
                    written.add(name)

                # settings.xml is always patched to Word 2016.
                _write_stream_setting(zin, zout)
                written.add("word/settings.xml")

                # Copy remaining template parts (theme, numbering, images...).
                _copy_template_extra_parts(zin, zout, written)

                # Replace document.xml with the merged body.
                with open(temp_docxml, "rb") as f:
                    zout.writestr("word/document.xml", f.read())

    finally:
        temp_body_file.unlink(missing_ok=True)
        temp_docxml.unlink(missing_ok=True)

    print(f"Merged .docx → {output_path}")
    return output_path


def merge_docx_files(source_dir, output_dir, use_python_docx=False):
    if use_python_docx and HAS_PYTHON_DOCX:
        return merge_docx_files_python_docx(source_dir, output_dir)
    return merge_docx_files_streaming(source_dir, output_dir)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main(argv=None):
    import argparse

    parser = argparse.ArgumentParser(description="Merge large TXT/DOCX chunks and inspect DOCX line counts.")
    parser.add_argument("source_dir", nargs="?", default=".",
                        help="Directory containing c1to700/c701to1k/ending files (default: cwd)")
    parser.add_argument("--inspect", nargs="+", metavar="DOCX",
                        help="Inspect DOCX paragraph/line counts instead of merging")
    parser.add_argument("--verify", action="store_true",
                        help="After merging, open merged.docx again with python-docx")
    parser.add_argument("--txt-only", action="store_true",
                        help="Merge TXT files only")
    parser.add_argument("--docx-only", action="store_true",
                        help="Merge DOCX files only")
    parser.add_argument("--use-python-docx", action="store_true",
                        help="Merge with python-docx instead of streaming "
                             "(small files only; loads whole DOCX files into RAM)")
    args = parser.parse_args(argv)

    if args.inspect:
        if not HAS_PYTHON_DOCX:
            print("[!!] --inspect requires python-docx.")
            return 2
        failed = False
        for path in args.inspect:
            try:
                inspect_docx(path)
            except Exception as e:
                print(f"[!!] Could not inspect {path}: {e}")
                failed = True
        return 1 if failed else 0

    if not HAS_PYTHON_DOCX:
        print("[i] python-docx not found. Install it with:")
        print("    pip install python-docx")
        print("    Alpine Linux: apk add py3-lxml && pip install python-docx")
        print("    (streaming DOCX merge does NOT require python-docx)")

    source_dir = Path(args.source_dir).resolve()
    output_dir = source_dir / "output"
    output_dir.mkdir(exist_ok=True)
    print(f"Source: {source_dir}")
    print(f"Output: {output_dir}")

    try:
        if args.txt_only:
            merge_txt_files(source_dir, output_dir)
        elif args.docx_only:
            check_docx_inputs(source_dir)
            merge_docx_files(source_dir, output_dir, use_python_docx=args.use_python_docx)
        else:
            merge_txt_files(source_dir, output_dir)
            check_docx_inputs(source_dir)
            merge_docx_files(source_dir, output_dir, use_python_docx=args.use_python_docx)
    except Exception as e:
        print(f"[!!] DOCX merge failed: {e}")
        if args.use_python_docx:
            print("    (--use-python-docx still failed.)")
        else:
            print("    For files with images/hyperlinks, retry with: "
                  "python tools/merge.py source --use-python-docx")
        return 1

    if args.verify:
        if not HAS_PYTHON_DOCX:
            print("[!!] --verify requires python-docx.")
        else:
            merged = output_dir / "merged.docx"
            if merged.exists():
                verify_docx(merged)

    print("Done.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
