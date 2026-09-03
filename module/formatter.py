from collections import Counter
from pathlib import Path

from .metadata import get_page_size, map_font_to_ttf


_MM_FONT = "YoeYar-One"
_EN_FONT = "Anonymous Pro"
_MM_HINTS = (
    "zawgyi", "myanmar", "amyanmar", "yoeyar", "pyidaungsu",
    "padauk", "notosansmyanmar", "tharlon", "sabae",
)
_EN_HINTS = ("times", "anonymous", "arial", "helvetica", "courier", "gautami", "roman")


def _is_myanmar_text(text: str) -> bool:
    return any("\u1000" <= c <= "\u109F" or "\uAA60" <= c <= "\uAA7F" for c in (text or ""))


# Internal plumbing keys that must not surface in the document header.
_INTERNAL_META_KEYS = {"page_count", "page_size", "font_map"}


def _flat_meta_items(metadata: dict):
    """Yield only scalar metadata entries for headers."""
    if not metadata:
        return
    for k, v in metadata.items():
        if isinstance(v, (dict, list, tuple)):
            continue
        if k in _INTERNAL_META_KEYS:
            continue
        yield k, v


def _build_meta_header(pdf_path, metadata, page_count):
    """Build the plain-text metadata header."""
    meta_lines = [f"# Source: {pdf_path}", f"# Pages: {page_count}", "#" * 50]
    for k, v in _flat_meta_items(metadata):
        meta_lines.insert(-1, f"# {k}: {v}")
    return "\n".join(meta_lines) + "\n\n"


def word_font_for(pdf_font: str | None, text: str = "", metadata: dict | None = None) -> str:
    """Map a PDF font family (or run text) to a Word-safe installed font name."""
    path = map_font_to_ttf(pdf_font or "", metadata) if pdf_font else None
    if path:
        if "Anonymous" in path:
            return _EN_FONT
        return _MM_FONT
    name = (pdf_font or "").lower()
    if any(h in name for h in _MM_HINTS):
        return _MM_FONT
    if any(h in name for h in _EN_HINTS):
        return _EN_FONT
    return _MM_FONT if _is_myanmar_text(text) else _EN_FONT


def page_separator(page_number: int) -> str:
    """The line that marks the start of a page in the plain-text output.

    TXT has no fonts, no margins and no page boxes, so a page is nothing more
    than a run of lines behind this marker.
    """
    return f"-------------------Page {page_number} -----------------\n"


def write_txt(all_texts, out_path, pdf_path, metadata):
    meta = _build_meta_header(pdf_path, metadata, len(all_texts))
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(meta)
        for idx, txt in enumerate(all_texts):
            f.write(page_separator(idx + 1))
            f.write(txt if txt else "")
            f.write("\n\n")
    print(f"[+] Saved TXT: {out_path}")


# ---------------------------------------------------------------------------
# Layout measurement (shared by the python-docx writer and the manual writer)
# ---------------------------------------------------------------------------

def _measure_layout(metadata, page_metadata):
    """Return the page geometry / line-pitch values used by the DOCX writers.

    Everything here is measured from the PDF itself: page box, text block
    edges, top/bottom margins and the real baseline pitch. Nothing assumes
    Word's 1" defaults, because these PDFs do not use them (the sample corpus
    sits at a 56.6 pt left edge with a 49 pt baseline pitch at 20 pt type).
    """
    width_pt, height_pt = get_page_size(metadata)
    # Word page size is in twips (1 pt = 20 twips).
    pg_w = max(1, int(round(width_pt * 20)))
    pg_h = max(1, int(round(height_pt * 20)))

    page_w, page_h = width_pt, height_pt

    # ── Text block edges, measured ─────────────────────────────────────
    # Left: the modal line start (body edge), not a percentile, so a few
    # indented lines cannot drag it. Right: a percentile of line ends,
    # clamped inside the page box -- glyph-advance rounding makes a small
    # tail of lines compute a `right` beyond the physical page.
    left_x = 72.0
    right_x = min(540.0, width_pt - 72.0)
    body_size = 12.0
    if page_metadata:
        xs, rr, szs = [], [], []
        for pg in page_metadata:
            for ln in (pg or {}).get("lines") or []:
                if not (ln.get("text") or "").strip():
                    continue
                if ln.get("x") is not None:
                    xs.append(round(float(ln["x"]), 1))
                r = ln.get("right")
                if r:
                    rr.append(float(r))
                if ln.get("size"):
                    szs.append(round(float(ln["size"]), 1))
        if xs:
            left_x = Counter(xs).most_common(1)[0][0]
        if szs:
            body_size = Counter(szs).most_common(1)[0][0]
        if rr:
            rr.sort()
            # `right` is x plus the sum of glyph advances. For Zawgyi the
            # stacked marks carry a nominal advance but render zero-width,
            # so summed widths overshoot the true inked extent (measured
            # ~11% long against the rendered page). Use the widest observed
            # line as the column requirement rather than a percentile, and
            # let the page box cap it: a column that is too NARROW makes
            # Word re-wrap and destroys the PDF's line breaks, while one a
            # little too wide changes nothing visually.
            est = rr[-1]
            # Leave a hairline gutter only; do not mirror the left margin,
            # which would clip the widest lines and re-wrap them.
            right_x = min(est, width_pt - 18.0)
        if right_x <= left_x:
            right_x = max(left_x + 72.0, width_pt - left_x)

    margin_left = max(0, int(round(left_x * 20)))

    # ── Measured rhythm (per paragraph, so mixed TOC/prose pages work) ──
    body_pitch = 0.0
    wide_rights = []
    right_left_pairs = []
    if page_metadata:
        all_gaps = []
        for pg in page_metadata:
            prev_y = None
            for ln in (pg or {}).get("lines") or []:
                if not (ln.get("text") or "").strip():
                    continue
                y = ln.get("y")
                if prev_y is not None and y is not None:
                    d = prev_y - y
                    # Only reject nonsense (column resets, overlapping runs).
                    # The real pitch is discovered below, not assumed to sit
                    # under some multiple of the font size.
                    if 0.3 < d < 0.5 * page_h:
                        all_gaps.append(d)
                prev_y = y
                r = ln.get("right") or 0
                w = r - (ln.get("x") or 0)
                right_left_pairs.append((r, w))
                if r:
                    wide_rights.append(r)
        if all_gaps:
            # The single-spaced baseline pitch is the MODE of the gap
            # distribution: the most repeated spacing is by definition
            # normal line advance. Median would be skewed by paragraph
            # breaks. Quantise to 0.5 pt so float noise clusters.
            hist = Counter(round(g * 2) / 2 for g in all_gaps)
            body_pitch = hist.most_common(1)[0][0]
    if body_pitch <= 0:
        body_pitch = 1.88 * 12
    pitch_tw = max(240, int(round(body_pitch * 20)))

    # A gap counts as a real blank line only when it clearly exceeds the
    # measured single-space pitch, not a font-size guess.
    para_gap = body_pitch * 1.45

    justified = False
    if wide_rights:
        full_thresh = 0.55 * ((right_x - left_x) or 1)
        rights = [r for r, w in right_left_pairs if w > full_thresh]
        if len(rights) >= 10:
            # Flush-right within half a character = justified body text.
            flush = sum(1 for r in rights if abs(r - right_x) <= 6.0)
            justified = flush / len(rights) > 0.55

    # Right margin follows the measured text block. Floor at 0 (not 0.75")
    # so a wide text block is not force-narrowed, which would re-wrap every
    # line and destroy the PDF's line breaks.
    margin_right = max(0, int(round((page_w - right_x) * 20)))

    # First-line paragraph indents: lines that start after a big gap
    # (paragraph start) but sit right of the body edge by a constant.
    first_line_tw = 0
    margin_top = 1440
    margin_bottom = 1440
    if page_metadata:
        starts = []
        tops, bottoms = [], []
        spans = []
        for pg in page_metadata:
            prev_y = None
            page_ys = []
            for ln in (pg or {}).get("lines") or []:
                if not (ln.get("text") or "").strip():
                    continue
                y = ln.get("y")
                if y is not None:
                    page_ys.append(y)
                if prev_y is not None and y is not None and (prev_y - y) > para_gap:
                    dx = (ln.get("x") or 0) - left_x
                    if 0.25 * 14 < dx < 3 * 14:
                        starts.append(dx)
                prev_y = y
            if page_ys:
                tops.append(max(page_ys))
                bottoms.append(min(page_ys))
                if len(page_ys) > 1:
                    spans.append(max(page_ys) - min(page_ys))
        if len(starts) >= 8:
            starts.sort()
            first_line_tw = int(round(starts[len(starts) // 2] * 20))
        # `y` is a BASELINE. The visual top of the text block sits one
        # ascent above it (~0.8 em for these faces), so the page margin is
        # measured to the glyph top, not to the baseline.
        ascent = body_size * 0.8
        if tops:
            tops.sort()
            first_baseline = tops[len(tops) // 2]
            margin_top = max(0, int(round((page_h - first_baseline - ascent) * 20)))
        if bottoms:
            bottoms.sort()
            last_baseline = bottoms[len(bottoms) // 2]
            descent = body_size * 0.2
            margin_bottom = max(0, int(round((last_baseline - descent) * 20)))

        # The text block must actually hold the tallest page. Word breaks a
        # page as soon as the next line does not fit, so a block even one
        # line short cascades every page onto two and doubles the document.
        #
        # The requirement is the measured first-to-last BASELINE SPAN, not
        # lines * body_pitch: a document can mix rhythms (c1-700 carries both
        # a 49 pt and a 26.5 pt pitch), and multiplying the busiest page's
        # line count by the single most common pitch overstates the need so
        # badly it goes past the paper (18 x 49 = 882 pt on a 792 pt page)
        # and pins both margins to the floor. Each page already carries its
        # own spacing, so the span is the honest number.
        if spans:
            spans.sort()
            need_pt = spans[-1] + ascent + body_size * 0.2
            need_tw = int(round(need_pt * 20))
            floor_tw = 288  # 0.2" -- keep the block off the paper edge
            slack = pg_h - margin_top - margin_bottom - need_tw
            if slack < 0:
                take = min(max(0, margin_bottom - floor_tw), -slack)
                margin_bottom -= take
                slack += take
            if slack < 0:
                take = min(max(0, margin_top - floor_tw), -slack)
                margin_top -= take

    return {
        "pg_w": pg_w,
        "pg_h": pg_h,
        "page_w": page_w,
        "page_h": page_h,
        "margin_left": margin_left,
        "margin_right": margin_right,
        "margin_top": margin_top,
        "margin_bottom": margin_bottom,
        "para_gap": para_gap,
        "pitch_tw": pitch_tw,
        "body_pitch": body_pitch,
        "justified": justified,
        "first_line_tw": first_line_tw,
        "left_x": left_x,
        "right_x": right_x,
    }


# ---------------------------------------------------------------------------
# Static OOXML package parts (shared by every DOCX writer)
# ---------------------------------------------------------------------------

def _content_types():
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        '<Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>'
        '<Override PartName="/word/settings.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.settings+xml"/>'
        '<Override PartName="/word/fontTable.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.fontTable+xml"/>'
        '<Override PartName="/word/webSettings.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.webSettings+xml"/>'
        '<Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>'
        '<Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>'
        "</Types>"
    )

def _root_rels():
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
        '<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>'
        '<Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>'
        "</Relationships>"
    )

def _document_rels():
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>'
        '<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/settings" Target="settings.xml"/>'
        '<Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/webSettings" Target="webSettings.xml"/>'
        '<Relationship Id="rId4" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/fontTable" Target="fontTable.xml"/>'
        "</Relationships>"
    )

def _styles_xml():
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:docDefaults>'
        '<w:rPrDefault><w:rPr>'
        '<w:rFonts w:ascii="Anonymous Pro" w:hAnsi="Anonymous Pro" w:cs="YoeYar-One" w:eastAsia="YoeYar-One"/>'
        '<w:sz w:val="24"/><w:szCs w:val="24"/>'
        '</w:rPr></w:rPrDefault>'
        '<w:pPrDefault><w:pPr><w:spacing w:line="240" w:lineRule="auto"/></w:pPr></w:pPrDefault>'
        '</w:docDefaults>'
        '<w:style w:type="paragraph" w:default="1" w:styleId="Normal"><w:name w:val="Normal"/><w:qFormat/></w:style>'
        '<w:style w:type="character" w:default="1" w:styleId="DefaultParagraphFont"><w:name w:val="Default Paragraph Font"/></w:style>'
        '<w:style w:type="table" w:default="1" w:styleId="TableNormal"><w:name w:val="Normal Table"/><w:qFormat/></w:style>'
        '<w:style w:type="paragraph" w:styleId="Title"><w:name w:val="Title"/><w:basedOn w:val="Normal"/><w:qFormat/>'
        '<w:pPr><w:spacing w:before="0" w:after="120"/><w:jc w:val="left"/></w:pPr>'
        '<w:rPr><w:b/><w:sz w:val="28"/><w:szCs w:val="28"/></w:rPr></w:style>'
        '<w:style w:type="paragraph" w:styleId="Heading1"><w:name w:val="heading 1"/><w:basedOn w:val="Normal"/><w:qFormat/>'
        '<w:pPr><w:keepNext/><w:spacing w:before="240" w:after="120"/><w:jc w:val="left"/></w:pPr>'
        '<w:rPr><w:b/><w:sz w:val="28"/><w:szCs w:val="28"/></w:rPr></w:style>'
        "</w:styles>"
    )

def _settings_xml():
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:settings xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:zoom w:percent="100"/>'
        "<w:defaultTabStop w:val=\"720\"/>"
        '<w:compat>'
        '<w:compatSetting w:name="compatibilityMode" w:uri="http://schemas.microsoft.com/office/word" w:val="16"/>'
        "</w:compat>"
        "</w:settings>"
    )

def _web_settings_xml():
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:webSettings xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "<w:optimizeForBrowser/>"
        "</w:webSettings>"
    )

def _font_table_xml():
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:fonts xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:font w:name="YoeYar-One"><w:charset w:val="00"/><w:family w:val="auto"/><w:pitch w:val="variable"/></w:font>'
        '<w:font w:name="Anonymous Pro"><w:charset w:val="00"/><w:family w:val="auto"/><w:pitch w:val="variable"/></w:font>'
        "</w:fonts>"
    )

def _core_xml():
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<cp:coreProperties xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" '
        'xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" '
        'xmlns:dcmitype="http://purl.org/dc/dcmitype/" xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">'
        '<dc:title>Myanmar PDF extraction</dc:title>'
        '<dc:creator>pdf-text-extractor</dc:creator>'
        '<cp:lastModifiedBy>pdf-text-extractor</cp:lastModifiedBy>'
        '<dcterms:modified xsi:type="dcterms:W3CDTF">2026-01-01T00:00:00Z</dcterms:modified>'
        "</cp:coreProperties>"
    )

def _app_xml():
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" '
        'xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">'
        "<Application>pdf-text-extractor</Application>"
        "<AppVersion>16.0000</AppVersion>"
        "</Properties>"
    )


def _static_docx_parts() -> dict:
    """Every part of the package except word/document.xml, as name -> xml."""
    return {
        "[Content_Types].xml": _content_types(),
        "_rels/.rels": _root_rels(),
        "word/_rels/document.xml.rels": _document_rels(),
        "word/styles.xml": _styles_xml(),
        "word/settings.xml": _settings_xml(),
        "word/webSettings.xml": _web_settings_xml(),
        "word/fontTable.xml": _font_table_xml(),
        "docProps/core.xml": _core_xml(),
        "docProps/app.xml": _app_xml(),
    }


# ---------------------------------------------------------------------------
# python-docx writer — generates a real Word 2016+/Office Open XML package
# (styles, settings, font table, doc props, document.xml.rels, ...).
# ---------------------------------------------------------------------------

def _set_pydocx_run_font(run, font_name, size_pt):
    """Apply a font to a python-docx run, including complex-script (w:cs)."""
    from docx.shared import Pt
    from docx.oxml.ns import qn

    size_pt = size_pt or 12
    run.font.name = font_name
    try:
        run.font.size = Pt(max(2, int(round(float(size_pt)))))
    except (TypeError, ValueError):
        run.font.size = Pt(12)

    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = rPr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), font_name)


def _pydocx_add_run_text(p, text, font_name, size_pt):
    """Add one run, converting literal \\t to real Word tab stops."""
    parts = (text or "").split("\t")
    for i, part in enumerate(parts):
        r = p.add_run(part)
        _set_pydocx_run_font(r, font_name, size_pt)
        if i < len(parts) - 1:
            tab_run = p.add_run()
            tab_run.add_tab()
            _set_pydocx_run_font(tab_run, font_name, size_pt)


def _set_pydocx_compat_mode(doc):
    """Tell Word 2016+ this is a modern (not legacy/compatibility) document.

    python-docx's default template declares compatibilityMode=14 (Word 2010),
    which makes Word 2016+ show the document in Compatibility Mode and call it
    an "old" document.  Bumping it to 16 declares the Word 2016 format.
    """
    from docx.oxml.ns import qn

    settings = doc.settings.element
    compat = settings.find(qn("w:compat"))
    if compat is None:
        from docx.oxml import OxmlElement
        compat = OxmlElement("w:compat")
        settings.append(compat)
    found = False
    for cs in compat.findall(qn("w:compatSetting")):
        if cs.get(qn("w:name")) == "compatibilityMode":
            cs.set(qn("w:val"), "16")
            found = True
            break
    if not found:
        from docx.oxml import OxmlElement
        cs = OxmlElement("w:compatSetting")
        cs.set(qn("w:name"), "compatibilityMode")
        cs.set(qn("w:uri"), "http://schemas.microsoft.com/office/word")
        cs.set(qn("w:val"), "16")
        compat.append(cs)


def _write_docx_python_docx(all_texts, out_path, pdf_path, metadata, page_metadata):
    """Legacy python-docx emitter. NO LONGER ON THE WRITE PATH.

    Kept because it is the reference the pre-streaming baselines were measured
    against, and removing it would delete a name other tooling may import.
    write_docx() no longer calls it: python-docx builds its whole XML tree in
    memory (301 MB for 2,294 pages, 927 MB for 6,870), which is exactly the
    cost StreamDocxWriter exists to remove.
    """
    from docx import Document
    from docx.shared import Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING

    layout = _measure_layout(metadata, page_metadata)
    doc = Document()
    _set_pydocx_compat_mode(doc)

    section = doc.sections[0]
    section.page_width = Twips(layout["pg_w"])
    section.page_height = Twips(layout["pg_h"])
    section.left_margin = Twips(layout["margin_left"])
    section.right_margin = Twips(layout["margin_right"])
    section.top_margin = Twips(1440)
    section.bottom_margin = Twips(1440)

    def style_par(par, style_name):
        try:
            par.style = doc.styles[style_name]
        except Exception:
            pass
        return par

    def new_paragraph(runs, space_tw=None, alignment=None,
                      left_tw=0, first_line_tw=0, style_name=None):
        p = doc.add_paragraph()
        if style_name:
            style_par(p, style_name)
        pf = p.paragraph_format
        if space_tw:
            pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
            pf.line_spacing = Twips(int(space_tw))
        if left_tw:
            pf.left_indent = Twips(int(left_tw))
        if first_line_tw:
            pf.first_line_indent = Twips(int(first_line_tw))
        if alignment is not None:
            p.alignment = alignment
        for text, font_name, size_pt in runs:
            _pydocx_add_run_text(p, text, font_name, size_pt)
        return p

    # ── Header / metadata ──────────────────────────────────────────────
    new_paragraph([(f"Source: {pdf_path}", _EN_FONT, 14)],
                  style_name="Title")
    new_paragraph([(f"Pages: {len(all_texts)}", _EN_FONT, 12)])
    for k, v in _flat_meta_items(metadata or {}):
        new_paragraph([(f"{k}: {v}", _EN_FONT, 11)])

    for idx, txt in enumerate(all_texts):
        # Page heading paragraph with a real page break.
        p = new_paragraph([("", _EN_FONT, 14)], style_name="Heading 1")
        r = p.add_run()
        r.add_break(WD_BREAK.PAGE)
        _set_pydocx_run_font(r, _EN_FONT, 14)
        _pydocx_add_run_text(p, f"Page {idx+1}", _EN_FONT, 14)

        page_meta = None
        if page_metadata and idx < len(page_metadata):
            page_meta = page_metadata[idx]

        lines = (page_meta or {}).get("lines") if page_meta else None
        if lines:
            prev_y = None
            for line in lines:
                y = line.get("y")
                size = line.get("size") or 12
                space_tw = layout["pitch_tw"]
                if prev_y is not None and y is not None:
                    d = prev_y - y  # y grows upward in PDF space
                    if 0.6 * size < d <= 2.3 * size:
                        space_tw = max(240, int(round(d * 20)))
                    elif d > 2.3 * size:
                        blank = round(d / layout["body_pitch"]) - 1
                        for _ in range(max(0, min(blank, 30))):
                            new_paragraph([("", _EN_FONT, size)])
                prev_y = y

                runs = line.get("runs") or [{
                    "text": line.get("text", ""),
                    "font": line.get("font"),
                    "size": size,
                }]
                parts = []
                for r in runs:
                    t = r.get("text") or ""
                    rsize = r.get("size") or size
                    font = word_font_for(r.get("font") or line.get("font"), t, metadata)
                    parts.append((t, font, rsize))
                if not parts:
                    parts.append(("", _EN_FONT, 12))

                # Paragraph properties: original line pitch, alignment.
                alignment = None
                lw = (line.get("right") or 0) - (line.get("x") or 0)
                cx = (line.get("x") or 0) + lw / 2
                if lw and abs(cx - layout["page_w"] / 2) <= 12 and lw < 0.6 * (layout["right_x"] - layout["left_x"]):
                    alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif layout["justified"] and lw > 0.55 * (layout["right_x"] - layout["left_x"]):
                    alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                left_tw = 0
                try:
                    left_tw = max(0, int(round((float(line.get("x")) - layout["left_x"]) * 20)))
                except (TypeError, ValueError):
                    pass

                new_paragraph(
                    parts,
                    space_tw=space_tw,
                    alignment=alignment,
                    left_tw=left_tw,
                    first_line_tw=layout["first_line_tw"],
                )
        else:
            for line in (txt or "").split("\n"):
                font = word_font_for(None, line, metadata)
                new_paragraph([(line, font, 12)])

    doc.save(out_path)
    return doc


# ---------------------------------------------------------------------------
# Manual / fallback OOXML writer (kept for environments without python-docx).
# It now writes the standard supporting parts as well, so Word and
# python-docx can open the result instead of treating it as a legacy file.
# ---------------------------------------------------------------------------

def _write_docx_manual(all_texts, out_path, pdf_path, metadata, page_metadata):
    """Legacy whole-document OOXML emitter. NO LONGER ON THE WRITE PATH.

    Superseded by StreamDocxWriter, which emits the same package but writes
    word/document.xml incrementally and positions every line absolutely
    (from its own x / gap / size) instead of relative to one document-wide
    page setup computed by _measure_layout. Kept for reference and because no
    name in this module is deleted.
    """
    import zipfile

    width_pt, height_pt = get_page_size(metadata)
    # Word page size is in twips (1 pt = 20 twips).
    pg_w = max(1, int(round(width_pt * 20)))
    pg_h = max(1, int(round(height_pt * 20)))

    with zipfile.ZipFile(out_path, "w", zipfile.ZIP_DEFLATED) as zf:

        zf.writestr("[Content_Types].xml", _content_types())
        zf.writestr("_rels/.rels", _root_rels())
        zf.writestr("word/_rels/document.xml.rels", _document_rels())
        zf.writestr("word/styles.xml", _styles_xml())
        zf.writestr("word/settings.xml", _settings_xml())
        zf.writestr("word/webSettings.xml", _web_settings_xml())
        zf.writestr("word/fontTable.xml", _font_table_xml())
        zf.writestr("docProps/core.xml", _core_xml())
        zf.writestr("docProps/app.xml", _app_xml())

        def _esc(s):
            return (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")

        def font_rpr(font_name, size_pt=12):
            sz = max(2, int(round(float(size_pt or 12) * 2)))  # half-points
            return (f"<w:rPr>"
                    f'<w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}" '
                    f'w:cs="{font_name}" w:eastAsia="{font_name}"/>'
                    f'<w:sz w:val="{sz}"/><w:szCs w:val="{sz}"/>'
                    f"</w:rPr>")

        def run(text, font_name, size_pt=12):
            # Literal \t (inserted by the parser at big column gaps) becomes a
            # real Word tab; everything else stays one xml:space=preserve run.
            parts = [_esc(p) for p in (text or "").split("\t")]
            inner = "<w:tab/>".join(
                f'<w:t xml:space="preserve">{p}</w:t>' for p in parts)
            return f"<w:r>{font_rpr(font_name, size_pt)}{inner}</w:r>"

        body = []
        body.append(f'<w:p><w:pPr><w:pStyle w:val="Title"/></w:pPr>'
                    f'{run(f"Source: {pdf_path}", _EN_FONT, 14)}</w:p>')
        body.append(f'<w:p>{run(f"Pages: {len(all_texts)}", _EN_FONT, 12)}</w:p>')

        for k, v in _flat_meta_items(metadata or {}):
            body.append(f"<w:p>{run(f'{k}: {v}', _EN_FONT, 11)}</w:p>")

        layout = _measure_layout(metadata, page_metadata)
        left_x = layout["left_x"]
        body_pitch = layout["body_pitch"]
        pitch_tw = layout["pitch_tw"]
        para_gap = layout["para_gap"]
        justified = layout["justified"]
        first_line_tw = layout["first_line_tw"]
        margin_right = layout["margin_right"]
        margin_left = layout["margin_left"]
        margin_top = layout["margin_top"]
        margin_bottom = layout["margin_bottom"]
        page_w = layout["page_w"]
        right_x = layout["right_x"]

        def indent_twips(x):
            try:
                return max(0, int(round((float(x) - left_x) * 20)))
            except (TypeError, ValueError):
                return 0

        for idx, txt in enumerate(all_texts):
            body.append(f'<w:p><w:pPr><w:pStyle w:val="Heading1"/>'
                        f'<w:spacing w:before="240" w:after="120"/></w:pPr>'
                        f'<w:r><w:br w:type="page"/></w:r>'
                        f'{run(f"Page {idx+1}", _EN_FONT, 14)}</w:p>')

            page_meta = None
            if page_metadata and idx < len(page_metadata):
                page_meta = page_metadata[idx]

            lines = (page_meta or {}).get("lines") if page_meta else None
            if lines:
                prev_y = None
                for line in lines:
                    y = line.get("y")
                    size = line.get("size") or 12
                    # Baseline pitch is measured (the mode of the gap
                    # distribution), so a gap only means "the author pressed
                    # Enter" when it exceeds that measured pitch -- comparing
                    # against a font-size multiple mislabels every single line
                    # whenever the document is leaded loosely.
                    space_tw = pitch_tw
                    if prev_y is not None and y is not None:
                        d = prev_y - y  # y grows upward in PDF space
                        if d <= para_gap:
                            space_tw = max(240, int(round(d * 20)))
                        else:
                            space_tw = max(240, int(round(body_pitch * 20)))
                            blank = int(round(d / body_pitch)) - 1
                            for _ in range(max(0, min(blank, 30))):
                                body.append(f'<w:p>{run("", _EN_FONT, size)}</w:p>')
                    prev_y = y

                    runs = line.get("runs") or [{
                        "text": line.get("text", ""),
                        "font": line.get("font"),
                        "size": size,
                    }]
                    parts = []
                    for r in runs:
                        t = r.get("text") or ""
                        rsize = r.get("size") or size
                        font = word_font_for(r.get("font") or line.get("font"), t, metadata)
                        parts.append(run(t, font, rsize))
                    if not parts:
                        parts.append(run("", _EN_FONT, 12))

                    # Paragraph properties: original line pitch, alignment.
                    ppr = [f'<w:spacing w:line="{space_tw}" w:lineRule="atLeast"/>']
                    ind = indent_twips(line.get("x"))
                    lw = (line.get("right") or 0) - (line.get("x") or 0)
                    cx = (line.get("x") or 0) + lw / 2
                    if lw and abs(cx - page_w / 2) <= 12 and lw < 0.6 * (right_x - left_x):
                        ppr.append('<w:jc w:val="center"/>')
                    elif justified and lw > 0.55 * (right_x - left_x):
                        ppr.append('<w:jc w:val="both"/>')
                    if ind or first_line_tw:
                        ppr.append(f'<w:ind w:left="{ind}" w:firstLine="{first_line_tw}"/>')
                    body.append(f'<w:p><w:pPr>{"".join(ppr)}</w:pPr>{"".join(parts)}</w:p>')
            else:
                for line in (txt or "").split("\n"):
                    font = word_font_for(None, line, metadata)
                    body.append(f"<w:p>{run(line, font, 12)}</w:p>")

        doc = (f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
               f'<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
               f'<w:body>{"".join(body)}'
               f'<w:sectPr><w:pgSz w:w="{pg_w}" w:h="{pg_h}"/>'
               f'<w:pgMar w:top="{margin_top}" w:right="{margin_right}" '
               f'w:bottom="{margin_bottom}" w:left="{margin_left}" '
               f'w:header="720" w:footer="720" w:gutter="0"/></w:sectPr>'
               f"</w:body></w:document>")
        zf.writestr("word/document.xml", doc)


def write_docx(all_texts, out_path, pdf_path, metadata, page_metadata=None):
    """Write a Word 2016+ compatible .docx.

    RESTRUCTURED: there is no longer a separate batch DOCX path. The streaming
    writer is the one and only emitter, so a .docx is byte-identical no matter
    which entry point produced it and no path can spike memory. This function
    keeps its name and signature — it just feeds the pages it was handed to
    StreamDocxWriter one at a time.

    python-docx cannot be used for this: it builds the whole XML tree in
    memory before saving, which is exactly the cost being removed. It stays a
    declared dependency because tooling still reads documents with it.
    """
    writer = StreamDocxWriter(out_path, pdf_path, len(all_texts))
    for idx, txt in enumerate(all_texts):
        layout = None
        if page_metadata and idx < len(page_metadata):
            layout = page_metadata[idx]
        writer.write_page(txt, layout)
    writer.close(metadata or {})


def write_output(all_texts, out_path, pdf_path, metadata=None, page_metadata=None):
    """Route to the correct writer based on file extension."""
    out_path = Path(out_path)
    out_ext = out_path.suffix.lower()
    if out_ext == ".txt":
        write_txt(all_texts, out_path, pdf_path, metadata or {})
    elif out_ext == ".docx":
        write_docx(all_texts, out_path, pdf_path, metadata or {}, page_metadata)
    else:
        raise ValueError(f"Unsupported output format: {out_ext}")
    print("[+] Done.")


# ---------------------------------------------------------------------------
# Streaming writers
#
# These accept ONE page at a time and write it out immediately, so the
# pipeline never holds more than a single page in memory.
#
# StreamDocxWriter used to be streaming in name only: it appended every page
# to two lists and handed the whole document to the batch writer at close().
# It now writes word/document.xml incrementally through
# zipfile.ZipFile.open(..., "w"), which deflates as it goes. Peak memory is
# one page of dicts plus the deflate buffer, flat regardless of page count.
#
# Because nothing is buffered there is also no document-wide pre-scan, so
# every position is written ABSOLUTELY (indent from the line's own x, spacing
# from its own gap to the previous baseline, its own size) instead of relative
# to one global page setup. That is both cheaper and more faithful: a document
# that mixes line rhythms is reproduced as it is, not averaged.
# ---------------------------------------------------------------------------

_PT_PER_TWIP = 20.0


def _esc_xml(s: str) -> str:
    """Escape text for an XML text node, dropping characters XML forbids."""
    s = (s or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    if any(ord(c) < 0x20 and c not in "\t\n\r" for c in s):
        s = "".join(c for c in s if ord(c) >= 0x20 or c in "\t\n\r")
    return s


def _run_xml(text: str, font_name: str, size_pt=12) -> str:
    """One <w:r> with its font and size. Literal \\t becomes a real Word tab."""
    sz = max(2, int(round(float(size_pt or 12) * 2)))  # half-points
    rpr = (f"<w:rPr>"
           f'<w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}" '
           f'w:cs="{font_name}" w:eastAsia="{font_name}"/>'
           f'<w:sz w:val="{sz}"/><w:szCs w:val="{sz}"/>'
           f"</w:rPr>")
    parts = [_esc_xml(p) for p in (text or "").split("\t")]
    inner = "<w:tab/>".join(
        f'<w:t xml:space="preserve">{p}</w:t>' for p in parts)
    return f"<w:r>{rpr}{inner}</w:r>"


def _sect_pr_xml(pg_w: int, pg_h: int, margins: dict, landscape: bool) -> str:
    """A <w:sectPr> for one run of same-shaped pages."""
    orient = ' w:orient="landscape"' if landscape else ""
    return (f'<w:sectPr><w:pgSz w:w="{pg_w}" w:h="{pg_h}"{orient}/>'
            f'<w:pgMar w:top="{margins["top"]}" w:right="{margins["right"]}" '
            f'w:bottom="{margins["bottom"]}" w:left="{margins["left"]}" '
            f'w:header="720" w:footer="720" w:gutter="0"/></w:sectPr>')


def page_shape(layout: dict | None, fallback=(612.0, 792.0)) -> tuple:
    """(width_pt, height_pt, rotation) of one page, in DISPLAY orientation.

    The mediabox is the unrotated paper; /Rotate 90 or 270 turns it, so the
    page the reader sees — and therefore the Word section — has its width and
    height swapped. 180 keeps them.
    """
    layout = layout or {}
    mb = layout.get("mediabox")
    if mb and len(mb) == 4:
        w = float(mb[2]) - float(mb[0])
        h = float(mb[3]) - float(mb[1])
    else:
        w, h = fallback
    if w <= 0 or h <= 0:
        w, h = fallback
    rotation = int(layout.get("rotation") or 0) % 360
    if rotation in (90, 270):
        w, h = h, w
    return (round(w, 3), round(h, 3), rotation)


def measure_page_margins(layout: dict | None, width_pt: float,
                         height_pt: float) -> dict:
    """Margins in twips for the section a page belongs to.

    Measured from the page's own lines: the left edge is the smallest line x,
    the top margin is measured to the glyph top (one ascent above the highest
    baseline) and the bottom to the descent below the lowest one. A page with
    no usable lines (an image page, a blank one) gets a symmetric default —
    nothing is laid out on it, so the value only has to be legal.
    """
    default = max(0, int(round(min(72.0, width_pt / 8.0) * _PT_PER_TWIP)))
    margins = {"left": default, "right": default,
               "top": default, "bottom": default}

    lines = [ln for ln in ((layout or {}).get("lines") or [])
             if (ln.get("text") or "").strip()]
    if not lines:
        return margins

    xs = [float(ln["x"]) for ln in lines if ln.get("x") is not None]
    ys = [float(ln["y"]) for ln in lines if ln.get("y") is not None]
    sizes = [float(ln["size"]) for ln in lines if ln.get("size")]
    body_size = min(sizes) if sizes else 12.0

    if xs:
        margins["left"] = max(0, int(round(min(xs) * _PT_PER_TWIP)))
    if ys:
        ascent = body_size * 0.8
        descent = body_size * 0.2
        top_pt = height_pt - max(ys) - ascent
        bottom_pt = min(ys) - descent
        margins["top"] = max(0, int(round(top_pt * _PT_PER_TWIP)))
        margins["bottom"] = max(0, int(round(bottom_pt * _PT_PER_TWIP)))

    # Never let the text block be narrower than the text: a too-narrow column
    # makes Word re-wrap and destroys the PDF's own line breaks.
    margins["right"] = max(0, int(round(18.0 * _PT_PER_TWIP)))
    # Keep a hairline of paper on every side and guarantee a positive block.
    max_side = int((width_pt * _PT_PER_TWIP) // 2) - 20
    margins["left"] = min(margins["left"], max(0, max_side))
    max_vert = int((height_pt * _PT_PER_TWIP) // 2) - 20
    margins["top"] = min(margins["top"], max(0, max_vert))
    margins["bottom"] = min(margins["bottom"], max(0, max_vert))
    return margins


def page_scale_factor(layout: dict | None, height_pt: float,
                      margins: dict) -> float:
    """Shrink factor (<= 1.0) that keeps one PDF page on one Word page.

    Invariant 3 is that the output has exactly one Word page per PDF page. A
    rotated page, or a page whose paper differs from the rest of the document,
    can carry a text block taller than the section it lands in; when that
    happens the whole page is scaled down uniformly (font size and spacing
    together) until it fits.

    The test is purely geometric, straight off the mediabox and the measured
    baselines: first-to-last baseline span plus one ascent and one descent
    against the usable height. It cannot fire on a page whose block already
    fits, which is why it is provably a no-op on documents that are uniform.
    Line WIDTH is deliberately not considered: `right` is a sum of glyph
    advances and overshoots the inked extent for stacked Myanmar marks, so
    using it would shrink fonts on perfectly good pages.
    """
    lines = [ln for ln in ((layout or {}).get("lines") or [])
             if (ln.get("text") or "").strip()]
    if len(lines) < 2:
        return 1.0
    ys = [float(ln["y"]) for ln in lines if ln.get("y") is not None]
    if len(ys) < 2:
        return 1.0
    sizes = [float(ln["size"]) for ln in lines if ln.get("size")]
    body_size = max(sizes) if sizes else 12.0

    need_pt = (max(ys) - min(ys)) + body_size * 0.8 + body_size * 0.2
    usable_pt = height_pt - (margins["top"] + margins["bottom"]) / _PT_PER_TWIP
    if usable_pt <= 0 or need_pt <= usable_pt:
        return 1.0
    # A little headroom so Word's own rounding cannot push the last line over.
    return max(0.15, (usable_pt / need_pt) * 0.98)


def render_page_body(layout: dict | None, text: str, margins: dict,
                     metadata: dict | None = None, scale: float = 1.0,
                     page_break: bool = True, sect_pr: str = "") -> str:
    """Render ONE page as OOXML paragraphs. No document-wide state is used.

    Every value comes from the page itself:
      * indent   -- the line's own x, relative to the section's left margin;
      * spacing  -- w:before is the line's own gap to the previous baseline
                    minus the natural line height, floored at 0, and w:line is
                    the natural line height (atLeast). Putting the whole gap in
                    w:before while also keeping a line height would count the
                    gap twice and overflow every page; with the max(0, ...)
                    form the baseline-to-baseline distance equals the measured
                    gap exactly;
      * size     -- the line's own font size.

    `sect_pr`, when given, is attached to the LAST paragraph of the page,
    which is how a section break is expressed mid-document.
    """
    out = []
    left_margin_pt = margins["left"] / _PT_PER_TWIP

    def indent_tw(x):
        try:
            return max(0, int(round((float(x) - left_margin_pt) * _PT_PER_TWIP * scale)))
        except (TypeError, ValueError):
            return 0

    lines = (layout or {}).get("lines") or []
    if not lines and text:
        lines = [{"text": t} for t in (text or "").split("\n")]

    prev_y = None
    for line in lines:
        y = line.get("y")
        size = (line.get("size") or 12) * scale
        natural_pt = size * 1.2  # the height the line occupies on its own

        before_tw = 0
        if prev_y is not None and y is not None:
            gap_pt = (float(prev_y) - float(y)) * scale
            if gap_pt > 0:
                before_tw = max(0, int(round((gap_pt - natural_pt) * _PT_PER_TWIP)))
        if y is not None:
            prev_y = y

        runs = line.get("runs") or [{
            "text": line.get("text", ""),
            "font": line.get("font"),
            "size": line.get("size") or 12,
        }]
        parts = []
        for r in runs:
            t = r.get("text") or ""
            rsize = (r.get("size") or line.get("size") or 12) * scale
            font = word_font_for(r.get("font") or line.get("font"), t, metadata)
            parts.append(_run_xml(t, font, rsize))
        if not parts:
            parts.append(_run_xml("", _EN_FONT, size))

        ppr = [f'<w:spacing w:before="{before_tw}" w:after="0" '
               f'w:line="{max(1, int(round(natural_pt * _PT_PER_TWIP)))}" '
               f'w:lineRule="atLeast"/>']
        ind = indent_tw(line.get("x"))
        if ind:
            ppr.append(f'<w:ind w:left="{ind}"/>')
        out.append(f'<w:p><w:pPr>{"".join(ppr)}</w:pPr>{"".join(parts)}</w:p>')

    if not out:
        # An empty page (an image page, or one with no extractable text) still
        # has to occupy exactly one Word page — invariant 3. It keeps an empty
        # <w:pPr> so it can still carry a page break and a section break.
        out.append("<w:p><w:pPr></w:pPr></w:p>")

    if page_break:
        # The break belongs at the START of the page, on its first paragraph.
        out[0] = out[0].replace("<w:pPr>", '<w:pPr><w:pageBreakBefore/>', 1)

    if sect_pr:
        # A mid-document section break lives inside the last paragraph's pPr.
        last = out[-1]
        if "<w:pPr>" in last:
            out[-1] = last.replace("</w:pPr>", sect_pr + "</w:pPr>", 1)
        else:
            out[-1] = last.replace("<w:p>", f"<w:p><w:pPr>{sect_pr}</w:pPr>", 1)
    return "".join(out)


class StreamTxtWriter:
    """Writes TXT incrementally, one page per call.

    TXT carries no geometry at all — there is no font to select and no page to
    lay out — so a page is just its lines in reading order behind a page
    separator.
    """

    def __init__(self, out_path, pdf_path, page_count, metadata=None):
        self.out_path = Path(out_path)
        self.pdf_path = pdf_path
        self.page_count = page_count
        self._idx = 0
        self._fh = open(self.out_path, "w", encoding="utf-8")
        # /Info and the page count are known before the walk starts, so the
        # header can go out immediately — no buffering needed.
        self._fh.write(_build_meta_header(pdf_path, metadata or {}, page_count))

    def write_page(self, text, layout=None):
        self._idx += 1
        self._fh.write(page_separator(self._idx))
        self._fh.write(text or "")
        self._fh.write("\n\n")

    def close(self, metadata=None):
        self._fh.close()
        print(f"[+] Saved TXT: {self.out_path}")


class StreamDocxWriter:
    """Writes a DOCX incrementally: one page in, one page of XML out.

    word/document.xml is opened once through zipfile's incremental write API
    and every page is deflated into it as it arrives, so nothing document-sized
    is ever held. Sections are handled with a ONE-PAGE LOOKAHEAD: a page is
    only flushed once the next one has been seen, which is exactly the
    information needed to decide whether a section break has to be attached to
    it. That costs one page of memory, not a pre-scan.
    """

    def __init__(self, out_path, pdf_path, page_count):
        import zipfile

        self.out_path = out_path
        self.pdf_path = pdf_path
        self.page_count = page_count
        self._zf = zipfile.ZipFile(out_path, "w", zipfile.ZIP_DEFLATED)
        for name, xml in _static_docx_parts().items():
            self._zf.writestr(name, xml)
        self._fh = self._zf.open("word/document.xml", "w")
        self._fh.write(
            b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            b'<w:document xmlns:w="http://schemas.openxmlformats.org/'
            b'wordprocessingml/2006/main"><w:body>')
        self._metadata = None
        self._pending = None       # the one-page lookahead buffer
        self._idx = 0
        # Section state: the shape of the current run of pages and the margins
        # measured from the page that started it. O(1) memory.
        self._sect_shape = None
        self._sect_margins = None
        self._last_sect_pr = None
        self._break_next = False   # a section break already moved us to a page

    # -- internals ----------------------------------------------------------

    def _emit(self, page, next_shape):
        """Write one buffered page, given the shape of the page after it."""
        text, layout = page
        shape = page_shape(layout)

        if self._sect_shape is None or shape != self._sect_shape:
            # First page of a new run of same-shaped pages: its margins become
            # the section's margins and are reused until the shape changes.
            self._sect_shape = shape
            self._sect_margins = measure_page_margins(layout, shape[0], shape[1])
        margins = self._sect_margins

        scale = page_scale_factor(layout, shape[1], margins)

        # Look ahead: if the next page is a different shape, this page's last
        # paragraph carries the section break.
        sect_pr = ""
        starts_new_section = next_shape is not None and next_shape != shape
        if starts_new_section:
            pg_w = max(1, int(round(shape[0] * _PT_PER_TWIP)))
            pg_h = max(1, int(round(shape[1] * _PT_PER_TWIP)))
            sect_pr = _sect_pr_xml(pg_w, pg_h, margins, shape[0] > shape[1])

        # A section break already begins a new page, so the page right after
        # one must NOT also carry an explicit page break or Word inserts a
        # blank page between them (invariant 3).
        page_break = self._idx > 0 and not self._break_next

        xml = render_page_body(layout, text, margins, self._metadata,
                               scale=scale, page_break=page_break,
                               sect_pr=sect_pr)
        self._fh.write(xml.encode("utf-8"))

        self._break_next = starts_new_section
        if starts_new_section:
            self._sect_shape = None
            self._sect_margins = None
        else:
            # Remember the geometry so the final body-level sectPr matches the
            # last run of pages.
            pg_w = max(1, int(round(shape[0] * _PT_PER_TWIP)))
            pg_h = max(1, int(round(shape[1] * _PT_PER_TWIP)))
            self._last_sect_pr = _sect_pr_xml(pg_w, pg_h, margins,
                                              shape[0] > shape[1])
        self._idx += 1

    # -- public API ---------------------------------------------------------

    def write_page(self, text, layout=None):
        shape = page_shape(layout)
        if self._pending is not None:
            self._emit(self._pending, shape)
        self._pending = (text, layout)

    def close(self, metadata=None):
        self._metadata = metadata or {}
        if self._pending is not None:
            self._emit(self._pending, None)
            self._pending = None
        if self._last_sect_pr is None:
            pg_w = max(1, int(round(612.0 * _PT_PER_TWIP)))
            pg_h = max(1, int(round(792.0 * _PT_PER_TWIP)))
            self._last_sect_pr = _sect_pr_xml(
                pg_w, pg_h,
                {"top": 1440, "right": 1440, "bottom": 1440, "left": 1440},
                False)
        self._fh.write(self._last_sect_pr.encode("utf-8"))
        self._fh.write(b"</w:body></w:document>")
        self._fh.close()
        self._zf.close()
        print(f"[+] Saved DOCX: {self.out_path} (streaming)")


def open_stream_writer(out_path, pdf_path, doc, enabled: bool = True):
    """Return the streaming writer for `out_path`.

    `enabled` is kept for signature compatibility with existing callers. There
    is no longer a batch alternative to fall back to: streaming is the only
    way a document is written, so both values return a writer.
    """
    ext = Path(out_path).suffix.lower()
    page_count = getattr(doc, "page_count", 0)
    doc_meta = getattr(doc, "meta", None) or {}
    if ext == ".txt":
        return StreamTxtWriter(out_path, pdf_path, page_count, doc_meta)
    if ext == ".docx":
        return StreamDocxWriter(out_path, pdf_path, page_count)
    raise ValueError(f"Unsupported output format: {ext}")
